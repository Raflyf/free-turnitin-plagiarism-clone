import fitz
import re
import logging
import datetime
import os
from typing import List, Tuple, Optional

logger = logging.getLogger(__name__)

# Pre-compiled Regex Patterns (Phase 3 #7)
RE_ZERO_WIDTH = re.compile(r'[\u200B-\u200D\uFEFF]')
RE_CYRILLIC_HOMOGLYPHS = re.compile(r'[асеорхуАСЕОРХУ]')
RE_SPACES = re.compile(r'\s+')
RE_HEADING_BAB1_LONG = re.compile(r'(?:BAB|CHAPTER)\s+(?:I|1)[\s:.\n]*(?:PENDAHULUAN|INTRODUCTION)\b')
RE_HEADING_BAB1_SHORT = re.compile(r'BAB\s+(?:I|1)\b')
RE_TOC_ENTRY_END = re.compile(r'[\s\.]*\d{1,3}\s*$')
RE_STRAIGHT_QUOTES = re.compile(r'"[^"]{1,500}"')
RE_SMART_QUOTES = re.compile(r'“[^”]{1,500}”')
RE_NEWLINES = re.compile(r'\n+')
RE_SENTENCE_SPLIT = re.compile(r'(?<=[.!?;])\s+')

def detect_manipulation(text: str, hidden_word_count: int = 0) -> List[str]:
    """Mendeteksi trik manipulasi dokumen untuk mencurangi mesin deteksi plagiarisme"""
    warnings = []
    # 1. Deteksi Zero-Width Characters (diselipkan antar huruf agar kata tidak terbaca)
    zero_width_chars = RE_ZERO_WIDTH.findall(text)
    if len(zero_width_chars) > 20:
        warnings.append("MANIPULASI TERDETEKSI: Ditemukan karakter tak terlihat (Zero-Width Space) yang digunakan untuk mengelabui sistem.")
    
    # 2. Deteksi huruf Cyrillic Homoglyphs (Huruf Rusia yang terlihat seperti huruf A, E, O latin)
    # Ini sangat umum digunakan untuk memutus N-Gram
    cyrillic_chars = RE_CYRILLIC_HOMOGLYPHS.findall(text)
    if len(cyrillic_chars) > 30:
        warnings.append("MANIPULASI TERDETEKSI: Ditemukan penggunaan huruf Cyrillic (Rusia) ilegal yang menyamar sebagai abjad Latin.")

    # 3. Deteksi teks tersembunyi (font mungil / warna putih) yang disuntik untuk
    # menggelembungkan jumlah kata sehingga persentase similarity turun. Kata semacam
    # ini sudah dibuang saat ekstraksi; di sini hanya memberi peringatan bila signifikan.
    if hidden_word_count > 30:
        warnings.append(f"MANIPULASI TERDETEKSI: Ditemukan ~{hidden_word_count} kata teks tersembunyi (font mungil/tak terlihat) yang disuntikkan untuk menurunkan persentase similarity.")

    return warnings

# Ambang teks tersembunyi: font < 4pt dianggap tak terbaca mata manusia
MIN_VISIBLE_FONT_SIZE = 4.0

def _extract_visible_text(doc):
    """Ekstrak teks, BUANG span dengan font mungil (< 4pt) yang tak terbaca mata.
    Mengembalikan (visible_text, hidden_word_count, any_dropped, hidden_spans).
    hidden_spans: list (page_index, (x0,y0,x1,y1)) untuk highlight di report PDF.

    PENTING: warna putih SENGAJA tidak dipakai sebagai sinyal, karena teks putih
    dipakai secara sah untuk label di atas kotak diagram (mis. "Raw Email",
    "Tokenizing") -> membuangnya = false positive + mengubah denominator dokumen bersih.
    Font mungil (< 4pt) adalah sinyal anti-cheat yang andal dan tak muncul di dokumen normal.

    Bila TIDAK ada span yang dibuang, caller memakai get_text() polos (verbatim) agar
    hasil bit-identik dengan sebelumnya -> skor dokumen bersih dijamin tidak berubah."""
    visible_parts = []
    hidden_word_count = 0
    any_dropped = False
    hidden_spans = []
    for page_index, page in enumerate(doc):
        pd = page.get_text("dict")
        for block in pd.get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    span_text = span.get("text", "")
                    if not span_text.strip():
                        continue
                    if span.get("size", 12.0) < MIN_VISIBLE_FONT_SIZE:
                        hidden_word_count += len(span_text.split())
                        any_dropped = True
                        bbox = span.get("bbox")
                        if bbox:
                            hidden_spans.append((page_index, tuple(bbox)))
                        continue
                    visible_parts.append(span_text)
                visible_parts.append(" ")
            visible_parts.append(" ")
    return "".join(visible_parts), hidden_word_count, any_dropped, hidden_spans

def extract_text_from_pdf(filepath: str, exclude_quotes: bool = True, exclude_biblio: bool = True, return_hidden: bool = False, fast_mode: bool = False, exclude_abstract: bool = True):
    """Extract text from PDF with robust error handling"""
    text = ""
    hidden_word_count = 0
    doc = None
    try:
        doc = fitz.open(filepath)
        # Deteksi teks font-mungil (anti-cheat). HANYA bila ada yang dibuang kita
        # pakai teks hasil span; jika tidak, pakai get_text() polos (verbatim) agar
        # dokumen bersih bit-identik -> skor tidak berubah. Robust: gagal -> get_text().
        try:
            if fast_mode:
                vis_text, hidden_word_count, any_dropped, hidden_spans = "", 0, False, []
            else:
                vis_text, hidden_word_count, any_dropped, hidden_spans = _extract_visible_text(doc)
        except Exception as vis_e:
            logger.warning("Failed to extract visible text, falling back to raw text: %s", vis_e)
            vis_text, hidden_word_count, any_dropped, hidden_spans = "", 0, False, []
        # Teks mentah (semua span, termasuk hidden) untuk skor "fooled"
        raw_text = ""
        for page in doc:
            raw_text += page.get_text() + " "

        if any_dropped and vis_text.strip():
            text = vis_text
        else:
            text = raw_text

        if not text.strip():
            raise Exception("PDF appears to be empty or contains only images")

    except Exception as e:
        raise Exception(f"Failed to extract PDF: {str(e)}")
    finally:
        if doc is not None:
            doc.close()

    manipulation_warnings = detect_manipulation(text, hidden_word_count)
    
    cleaned_text = clean_text(text, exclude_quotes, exclude_biblio, exclude_abstract)
    
    # Bersihkan Zero-width chars dari teks agar tetap bisa di-cek similarity-nya
    cleaned_text = RE_ZERO_WIDTH.sub('', cleaned_text)
    # Normalkan huruf Cyrillic kembali ke Latin agar usahanya sia-sia
    cyrillic_to_latin = str.maketrans('асеорхуАСЕОРХУ', 'aceopxyACEOPXY')
    cleaned_text = cleaned_text.translate(cyrillic_to_latin)

    if return_hidden:
        # Bersihkan raw_text dengan cara yang sama
        raw_cleaned = clean_text(raw_text, exclude_quotes, exclude_biblio, exclude_abstract)
        raw_cleaned = RE_ZERO_WIDTH.sub('', raw_cleaned)
        raw_cleaned = raw_cleaned.translate(cyrillic_to_latin)
        return cleaned_text, manipulation_warnings, raw_cleaned, hidden_spans

    return cleaned_text, manipulation_warnings

def _extract_visible_docx(doc):
    """
    Ekstrak teks dari DOCX, BUANG bagian (run) yang memiliki font mungil (< 4pt),
    berwarna putih, atau diatur tersembunyi (hidden).
    """
    visible_parts = []
    raw_parts = []
    hidden_word_count = 0
    any_dropped = False
    
    def process_runs(runs):
        nonlocal hidden_word_count, any_dropped
        for run in runs:
            text = run.text
            if not text.strip():
                raw_parts.append(text)
                visible_parts.append(text)
                continue
                
            is_hidden = False
            # 1. Cek properti font 'hidden'
            if run.font.hidden:
                is_hidden = True
            # 2. Cek font size sangat kecil (< 4pt)
            elif getattr(run.font.size, 'pt', None) is not None and getattr(run.font.size, 'pt', None) < 4:
                is_hidden = True
            # 3. Cek warna putih
            elif run.font.color and run.font.color.rgb and str(run.font.color.rgb) == "FFFFFF":
                is_hidden = True
                
            raw_parts.append(text)
            
            if is_hidden:
                hidden_word_count += len(text.split())
                any_dropped = True
            else:
                visible_parts.append(text)
                
    for p in doc.paragraphs:
        process_runs(p.runs)
        raw_parts.append("\n")
        visible_parts.append("\n")
        
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_runs(p.runs)
                    raw_parts.append("\n")
                    visible_parts.append("\n")
                    
    return "".join(visible_parts), hidden_word_count, any_dropped, "".join(raw_parts)


def extract_text_from_docx(docx_path: str, exclude_quotes: bool = True, exclude_biblio: bool = True, return_hidden: bool = False, fast_mode: bool = False, exclude_abstract: bool = True):
    """Extract text from .docx (Word). Mendeteksi trik manipulasi font."""
    from docx import Document
    doc = Document(docx_path)
    
    vis_text, hidden_word_count, any_dropped, raw_text = _extract_visible_docx(doc)
    
    if not raw_text.strip():
        raise Exception("DOCX appears to be empty")
        
    # Gunakan text asli (raw) jika fast_mode atau tidak ada manipulasi agar skor jujur 100% konsisten
    if any_dropped and vis_text.strip() and not fast_mode:
        text = vis_text
    else:
        text = raw_text

    manipulation_warnings = detect_manipulation(text, hidden_word_count)
    cleaned_text = clean_text(text, exclude_quotes, exclude_biblio, exclude_abstract)
    cleaned_text = RE_ZERO_WIDTH.sub('', cleaned_text)
    cyrillic_to_latin = str.maketrans('асеорхуАСЕОРХУ', 'aceopxyACEOPXY')
    cleaned_text = cleaned_text.translate(cyrillic_to_latin)
    
    if return_hidden:
        raw_cleaned = clean_text(raw_text, exclude_quotes, exclude_biblio, exclude_abstract)
        raw_cleaned = RE_ZERO_WIDTH.sub('', raw_cleaned)
        raw_cleaned = raw_cleaned.translate(cyrillic_to_latin)
        return cleaned_text, manipulation_warnings, raw_cleaned, []
    return cleaned_text, manipulation_warnings


def extract_text_auto(filepath, exclude_quotes=True, exclude_biblio=True, return_hidden=False, fast_mode=False, exclude_abstract=True):
    """Deteksi ekstensi lalu ekstrak (.pdf/.docx/.txt). Satu pintu untuk semua format."""
    low = filepath.lower()
    if low.endswith(".docx") or low.endswith(".doc"):
        return extract_text_from_docx(filepath, exclude_quotes, exclude_biblio, return_hidden=return_hidden, exclude_abstract=exclude_abstract)
    if low.endswith(".txt"):
        res = extract_text_from_txt(filepath)
        if return_hidden: return res, [], res, []
        return res, []
    return extract_text_from_pdf(filepath, exclude_quotes, exclude_biblio, return_hidden=return_hidden, fast_mode=fast_mode, exclude_abstract=exclude_abstract)


def extract_text_from_txt(txt_path):
    """Extract text from TXT with automatic encoding detection"""
    try:
        import chardet
        
        # Detect encoding
        with open(txt_path, 'rb') as f:
            raw_data = f.read()
        
        if not raw_data:
            raise Exception("File is empty")
        
        detected = chardet.detect(raw_data)
        encoding = detected['encoding'] or 'utf-8'
        
        # Try detected encoding with fallbacks
        encodings_to_try = [encoding, 'utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for enc in encodings_to_try:
            try:
                return raw_data.decode(enc)
            except (UnicodeDecodeError, AttributeError):
                continue
        
        # Last resort: decode with errors='replace'
        return raw_data.decode('utf-8', errors='replace')
        
    except Exception as e:
        raise Exception(f"Failed to extract TXT: {str(e)}")

def clean_text(text, exclude_quotes=True, exclude_biblio=True, exclude_abstract=True):
    text = RE_SPACES.sub(' ', text).strip()

    # [1] Exclude Front Matter (Cover, Pengesahan, Daftar Isi) - Standar Industri
    upper_text = text.upper()

    # Pola heading bab asli: "BAB I" / "BAB 1" diikuti KONTEN nyata (PENDAHULUAN).
    # Regex ini lebih presisi karena mencegah Lembar Konsultasi / Abstrak yang sekadar menyebut "BAB I".
    chosen_idx = -1
    for m in RE_HEADING_BAB1_LONG.finditer(upper_text):
        idx = m.start()
        # Ambil 40 karakter setelah match untuk cek apakah ini entri daftar isi
        tail = text[m.end():m.end() + 40]
        # Entri daftar isi: didominasi titik-titik atau langsung angka halaman
        dot_ratio = tail.count('.') / max(len(tail), 1)
        is_toc_entry = dot_ratio > 0.3 or bool(RE_TOC_ENTRY_END.match(tail[:15]))
        if not is_toc_entry:
            chosen_idx = idx
            break

    # Fallback: jika regex spesifik gagal (sangat jarang), gunakan regex lama
    if chosen_idx == -1:
        for m in RE_HEADING_BAB1_SHORT.finditer(upper_text):
            idx = m.start()
            tail = text[m.end():m.end() + 40]
            dot_ratio = tail.count('.') / max(len(tail), 1)
            is_toc_entry = dot_ratio > 0.3 or bool(RE_TOC_ENTRY_END.match(tail[:15]))
            if not is_toc_entry:
                chosen_idx = idx
                break

    if chosen_idx != -1 and chosen_idx < len(text) * 0.4:
        text = text[chosen_idx:]
    elif exclude_abstract:
        # Jika BAB 1 tidak ditemukan (chosen_idx = -1), kita cari bagian Abstrak dan hapus secara spesifik (untuk tipe jurnal).
        # Cari dari awal. Jika ada ABSTRAK / ABSTRACT, cari kata kunci akhir atau paragraf selanjutnya.
        m_abs = re.search(r'\b(?:ABSTRAK|ABSTRACT)\b', upper_text)
        if m_abs and m_abs.start() < len(text) * 0.3:
            # Cari akhir abstrak (bisa berupa "KATA KUNCI:", "KEYWORDS:", atau "1. PENDAHULUAN" / "1. INTRODUCTION")
            start_abs = m_abs.start()
            end_abs = -1
            m_end = re.search(r'\b(?:KATA KUNCI|KEYWORDS|PENDAHULUAN|INTRODUCTION)\b', upper_text[start_abs + 10:])
            if m_end:
                # Potong teks dari awal sampai KATA KUNCI (kita asumsikan setelah kata kunci adalah teks asli yang dicek)
                # Atau jika itu Pendahuluan, potong sampai pendahuluan
                # Lebih aman: hapus blok abstraknya.
                end_abs = start_abs + 10 + m_end.end()
            
            if end_abs != -1 and end_abs - start_abs < 3000:
                text = text[:start_abs] + text[end_abs:]


    # [2] Exclude Bibliography
    if exclude_biblio:
        last_idx = max(text.upper().rfind('DAFTAR PUSTAKA'), text.upper().rfind('REFERENCES'))
        if last_idx > len(text) * 0.5:
            text = text[:last_idx]

    # [3] Exclude Quotes
    if exclude_quotes:
        # Hapus kutipan dengan straight quotes (maks 500 karakter agar tidak menghapus 1 bab jika ada quote yg tidak tertutup)
        text = RE_STRAIGHT_QUOTES.sub('', text)
        # Hapus kutipan dengan smart quotes
        text = RE_SMART_QUOTES.sub('', text)

    return text

def get_sentences(text):
    text = RE_NEWLINES.sub('. ', text)
    sentences = RE_SENTENCE_SPLIT.split(text)
    return [s.strip() for s in sentences if len(s.split()) >= 5]